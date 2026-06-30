from __future__ import annotations

import re
from datetime import date
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image


BASE_DIR = Path(__file__).resolve().parent
PROJECT_ROOT = BASE_DIR.parent
OUTPUTS = {
    "onboarding": BASE_DIR / "Manual_Onboarding_ERP_OS.docx",
    "complete": BASE_DIR / "Manual_Completo_Utilizador_ERP_OS.docx",
}


def screenshot_files() -> list[Path]:
    return sorted(BASE_DIR.glob("*.png"), key=lambda p: p.name)


def set_cell_text(cell, text: str, bold: bool = False) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    run = paragraph.add_run(text)
    run.bold = bold


def add_toc(paragraph) -> None:
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'

    fld_sep = OxmlElement("w:fldChar")
    fld_sep.set(qn("w:fldCharType"), "separate")

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_sep)
    run._r.append(fld_end)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.add_run("Página ")
    run = paragraph.add_run()
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_end)


def configure_document(doc: Document, subtitle: str) -> None:
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(1.8)
    section.bottom_margin = Cm(1.8)
    section.left_margin = Cm(1.8)
    section.right_margin = Cm(1.8)

    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10.5)

    for style_name, size, color in [
        ("Title", 24, "17365D"),
        ("Heading 1", 18, "17365D"),
        ("Heading 2", 14, "1F4E79"),
        ("Heading 3", 12, "2F5597"),
    ]:
        style = styles[style_name]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    header = section.header.paragraphs[0]
    header.text = subtitle
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    header.runs[0].font.size = Pt(8)
    header.runs[0].font.color.rgb = RGBColor(100, 100, 100)


def add_cover(doc: Document, title: str, subtitle: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.space_after = Pt(36)
    run = p.add_run("ERP_OS / Junta Operacional")
    run.bold = True
    run.font.name = "Arial"
    run.font.size = Pt(14)
    run.font.color.rgb = RGBColor(31, 78, 121)

    title_p = doc.add_paragraph()
    title_p.style = "Title"
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_p.add_run(title)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(subtitle)
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(80, 80, 80)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    meta.add_run(f"Versão 1.0 | {date.today().strftime('%d/%m/%Y')}")

    doc.add_paragraph()
    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    note.add_run(
        "Documento editável em Word, preparado a partir das funcionalidades programadas "
        "e das capturas existentes na pasta Manual_ERP_OS."
    )

    doc.add_page_break()
    doc.add_heading("Índice", level=1)
    add_toc(doc.add_paragraph())
    doc.add_paragraph(
        "Nota: no Microsoft Word, clique com o botão direito sobre o índice e escolha "
        "'Atualizar campo' para atualizar números de página após edições."
    )
    doc.add_page_break()


def clean_inline(text: str) -> str:
    text = re.sub(r"\*\*(.*?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]*)`", r"\1", text)
    return text.strip()


def add_rich_paragraph(doc: Document, text: str, style: str | None = None):
    p = doc.add_paragraph(style=style)
    parts = re.split(r"(\*\*.*?\*\*)", text)
    for part in parts:
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            run = p.add_run(part[2:-2])
            run.bold = True
        else:
            p.add_run(part)
    return p


def image_size_for_page(path: Path, max_width_in: float = 6.25, max_height_in: float = 4.35) -> tuple[float, float]:
    with Image.open(path) as img:
        width, height = img.size
    ratio = min(max_width_in / width, max_height_in / height)
    return width * ratio, height * ratio


def add_screenshot(doc: Document, image_path: Path, caption: str) -> None:
    width, height = image_size_for_page(image_path)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(width), height=Inches(height))

    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = cap.add_run(caption)
    r.italic = True
    r.font.size = Pt(8.5)
    r.font.color.rgb = RGBColor(90, 90, 90)


def add_app_functionality_summary(doc: Document) -> None:
    doc.add_heading("Mapa funcional validado", level=1)
    intro = (
        "Este resumo cruza o conteúdo do manual com as funcionalidades programadas no projeto. "
        "Serve como orientação rápida para administradores funcionais e formadores."
    )
    doc.add_paragraph(intro)

    rows = [
        ("Autenticação e perfil", "Login, logout, perfil do utilizador e atualização de dados pessoais."),
        ("Pedidos", "Criação, consulta, comentários, anexos, encaminhamento, atribuição, cancelamento e geração de tarefas."),
        ("Tarefas", "Criação, edição, atribuição, checklist, validação, materiais associados e recorrência."),
        ("Equipas", "Gestão de equipas e membros."),
        ("Agenda e reservas", "Eventos, participantes, reservas, aprovação/rejeição e gestão de espaços."),
        ("Diretório", "Pessoas, entidades, fotos, contas de acesso e ausências associadas a pessoas."),
        ("Recursos humanos", "Funcionários, detalhe, edição e registo de ausências."),
        ("Inventário", "Catálogo, stock, movimentos, categorias, empréstimos, devoluções, requisições e alocações."),
        ("Documentos e atas", "Documentos, downloads, aprovação, pedido de aprovação, atas e visibilidade."),
        ("Planeamento", "Planos operacionais, agenda e requisições associadas."),
        ("Chat", "Conversas, mensagens, não lidas, push, apagar mensagens e conversão em tarefa ou pedido."),
        ("Relatórios", "Página de acompanhamento e análise operacional."),
        ("Configurações", "Instituição, áreas, utilizadores, perfis, permissões, delegações e tipos de pessoa."),
        ("Notificações", "Listagem, recentes, contador de não lidas e marcação como lida."),
    ]
    table = doc.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    set_cell_text(table.rows[0].cells[0], "Área", True)
    set_cell_text(table.rows[0].cells[1], "Funcionalidades", True)
    for area, funcs in rows:
        cells = table.add_row().cells
        set_cell_text(cells[0], area, True)
        set_cell_text(cells[1], funcs)


def render_markdown(doc: Document, markdown: str, images: list[Path], use_all_images: bool) -> set[Path]:
    used: set[Path] = set()
    image_index = 0
    in_code = False
    pending_blank = False

    lines = markdown.splitlines()
    for raw_line in lines:
        line = raw_line.rstrip()
        stripped = line.strip()

        if stripped.startswith("```"):
            in_code = not in_code
            continue
        if in_code:
            p = doc.add_paragraph(stripped)
            p.style = "Intense Quote"
            continue
        if not stripped:
            pending_blank = True
            continue
        if stripped.startswith("# "):
            doc.add_heading(clean_inline(stripped[2:]), level=1)
            pending_blank = False
            continue
        if stripped.startswith("## "):
            doc.add_heading(clean_inline(stripped[3:]), level=1)
            pending_blank = False
            continue
        if stripped.startswith("### "):
            doc.add_heading(clean_inline(stripped[4:]), level=2)
            pending_blank = False
            continue
        if stripped.startswith(">"):
            p = add_rich_paragraph(doc, clean_inline(stripped.lstrip("> ")), style="Intense Quote")
            p.paragraph_format.left_indent = Cm(0.5)
            continue

        image_hint = stripped.lower().startswith("imagem recomendada") or stripped.lower().startswith("imagens recomendadas")
        if image_hint:
            if image_index < len(images):
                count = 2 if use_all_images and stripped.lower().startswith("imagens recomendadas") else 1
                for _ in range(count):
                    if image_index >= len(images):
                        break
                    image_path = images[image_index]
                    used.add(image_path)
                    add_screenshot(
                        doc,
                        image_path,
                        f"Figura {len(used)} - Captura da aplicação: {image_path.stem}",
                    )
                    image_index += 1
            continue

        numbered = re.match(r"^(\d+)\.\s+(.*)$", stripped)
        if numbered:
            doc.add_paragraph(clean_inline(numbered.group(2)), style="List Number")
            continue
        if stripped.startswith("- "):
            doc.add_paragraph(clean_inline(stripped[2:]), style="List Bullet")
            continue

        if pending_blank and len(stripped) < 80 and stripped.endswith(":"):
            p = add_rich_paragraph(doc, clean_inline(stripped))
            p.runs[0].bold = True
        else:
            add_rich_paragraph(doc, stripped)
        pending_blank = False

    if use_all_images:
        remaining = [img for img in images if img not in used]
        if remaining:
            doc.add_page_break()
            doc.add_heading("Anexo: Capturas da aplicação", level=1)
            doc.add_paragraph(
                "As imagens seguintes complementam os procedimentos descritos no manual e ficam "
                "incorporadas no documento para consulta e edição posterior."
            )
            for img in remaining:
                used.add(img)
                add_screenshot(doc, img, f"Figura {len(used)} - Captura da aplicação: {img.stem}")

    return used


def build_document(kind: str, title: str, subtitle: str, markdown_file: Path, images: list[Path], use_all_images: bool) -> None:
    doc = Document()
    configure_document(doc, subtitle)
    add_cover(doc, title, subtitle)
    add_app_functionality_summary(doc)
    doc.add_page_break()

    markdown = markdown_file.read_text(encoding="utf-8")
    render_markdown(doc, markdown, images, use_all_images=use_all_images)

    doc.save(OUTPUTS[kind])


def main() -> None:
    images = screenshot_files()
    if not images:
        raise SystemExit("Não foram encontradas imagens PNG em Manual_ERP_OS.")

    build_document(
        "onboarding",
        "Manual de Onboarding",
        "Guia rápido para novos utilizadores",
        BASE_DIR / "Onboarding_Rapido.md",
        images[:18],
        use_all_images=False,
    )
    build_document(
        "complete",
        "Manual Completo de Utilizador",
        "Procedimentos completos de utilização da aplicação",
        BASE_DIR / "Manual_Completo_Utilizador.md",
        images,
        use_all_images=True,
    )

    print("Documentos criados:")
    for output in OUTPUTS.values():
        print(f"- {output}")


if __name__ == "__main__":
    main()
